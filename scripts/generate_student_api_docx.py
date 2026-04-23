from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def zh(text: str) -> str:
    return text.encode("ascii").decode("unicode_escape")


def set_run_font(run, name: str = "Microsoft YaHei", size: float | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)


def set_cell_text(cell, text: str) -> None:
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            set_run_font(run, size=10)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr_cells[idx], header)
        shade_cell(hdr_cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph("")


def build_document(output_path: Path) -> None:
    doc = Document()

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Title"].font.size = Pt(20)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(12.5)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(title.add_run(zh("\\u5b66\\u751f\\u7aef\\u63a5\\u53e3\\u6587\\u6863")))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run(zh("\\u57fa\\u4e8e\\u5f53\\u524d\\u9879\\u76ee\\u4ee3\\u7801\\u6574\\u7406")), size=10.5)

    doc.add_heading(zh("\\u57fa\\u7840\\u4fe1\\u606f"), level=1)
    for item in [
        zh("\\u540e\\u7aef\\u6846\\u67b6\\uff1aFastAPI"),
        zh("\\u5f53\\u524d\\u5b66\\u751f\\u7aef\\u5b9e\\u9645\\u63a5\\u53e3\\u524d\\u7f00\\uff1a/student-api"),
        zh("\\u524d\\u7aef\\u5b9e\\u9645 baseURL\\uff1aVITE_STUDENT_API_BASE || /student-api"),
        zh("\\u65e7\\u63a5\\u53e3\\u524d\\u7f00\\uff1a/api/v1"),
        zh("\\u6587\\u6863\\u5730\\u5740\\uff1ahttp://127.0.0.1:3001/docs"),
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading(zh("\\u901a\\u7528\\u7ea6\\u5b9a"), level=1)
    for item in [
        zh("\\u5b66\\u751f\\u7aef\\u6240\\u6709 POST \\u63a5\\u53e3\\u90fd\\u4f1a\\u81ea\\u52a8\\u9644\\u52a0\\u7b7e\\u540d\\u5b57\\u6bb5\\uff1atime\\u3001enc\\u3002"),
        zh("\\u5b66\\u751f\\u7aef\\u901a\\u5e38\\u8fd8\\u4f1a\\u5e26\\u8bf7\\u6c42\\u5934\\uff1aAuthorization: Bearer <token>\\u3001X-Platform-Token: <token>\\u3002"),
        zh("\\u5b66\\u751f\\u7aef\\u7edf\\u4e00\\u54cd\\u5e94\\u7ed3\\u6784\\u5982\\u4e0b\\u3002"),
    ]:
        doc.add_paragraph(item, style="List Bullet")

    for line in [
        "{",
        '  "code": 200,',
        '  "msg": "success",',
        '  "data": {},',
        '  "requestId": "req-xxxxxxxxxxxx"',
        "}",
    ]:
        p = doc.add_paragraph()
        set_run_font(p.add_run(line), name="Consolas", size=9.5)

    doc.add_heading(zh("\\u5b66\\u751f\\u7aef\\u5b9e\\u9645\\u63a5\\u53e3"), level=1)
    add_table(
        doc,
        [
            zh("\\u63a5\\u53e3"),
            zh("\\u65b9\\u6cd5"),
            zh("\\u4f5c\\u7528"),
            zh("\\u4e3b\\u8981\\u8bf7\\u6c42\\u53c2\\u6570"),
            zh("\\u4e3b\\u8981\\u8fd4\\u56de\\u6570\\u636e"),
        ],
        [
            ["/student-api/auth/verify", "POST", zh("\\u5b66\\u751f\\u514d\\u767b\\u6821\\u9a8c"), "token, platform, clientType, time, enc", "student, lessons"],
            ["/student-api/api/v1/getStudentLessonList", "POST", zh("\\u83b7\\u53d6\\u5b66\\u751f\\u8bfe\\u7a0b\\u5217\\u8868"), "studentId, token, time, enc", "lessons"],
            ["/student-api/api/v1/recentChapters/list", "POST", zh("\\u83b7\\u53d6\\u6700\\u8fd1\\u5b66\\u4e60\\u8bb0\\u5f55"), "studentId, limit, time, enc", "items"],
            ["/student-api/api/v1/recentChapters/save", "POST", zh("\\u4fdd\\u5b58\\u6700\\u8fd1\\u5b66\\u4e60\\u7ae0\\u8282"), "studentId, lessonId, sectionId, pageNo, time, enc", zh("\\u6700\\u8fd1\\u5b66\\u4e60\\u8bb0\\u5f55\\u5bf9\\u8c61")],
            ["/student-api/api/v1/progress/get", "POST", zh("\\u83b7\\u53d6\\u8bfe\\u7a0b\\u5b66\\u4e60\\u8fdb\\u5ea6"), "studentId, lessonId, time, enc", zh("\\u8fdb\\u5ea6\\u72b6\\u6001\\u5bf9\\u8c61")],
            ["/student-api/api/v1/progress/track", "POST", zh("\\u4e0a\\u62a5\\u5b66\\u4e60\\u8fdb\\u5ea6"), "lessonId, anchorId, anchorTitle, pageNo, currentTime, progressPercent, understandingLevel, weakPoints, time, enc", zh("\\u66f4\\u65b0\\u540e\\u7684\\u8fdb\\u5ea6\\u5bf9\\u8c61")],
            ["/student-api/api/v1/progress/page/read", "POST", zh("\\u6807\\u8bb0\\u67d0\\u9875\\u5df2\\u8bfb"), "studentId, lessonId, sectionId, lessonPageId, pageNo, time, enc", zh("\\u9875\\u9762\\u9605\\u8bfb\\u7ed3\\u679c\\u3001\\u7ae0\\u8282\\u8fdb\\u5ea6")],
            ["/student-api/api/v1/progress/adjust", "POST", zh("\\u8c03\\u6574\\u5b66\\u4e60\\u8282\\u594f\\u5efa\\u8bae"), "studentId, lessonId, anchorId, anchorTitle, pageNo, action, time, enc", "advice " + zh("\\u7b49\\u5efa\\u8bae\\u5185\\u5bb9")],
            ["/student-api/api/v1/lesson/play", "POST", zh("\\u83b7\\u53d6\\u8bfe\\u7a0b\\u64ad\\u653e\\u9875\\u6570\\u636e"), "studentId, lessonId, anchorId, time, enc", zh("\\u8bfe\\u7a0b\\u64ad\\u653e\\u6570\\u636e")],
            ["/student-api/api/v1/lesson/section/detail", "POST", zh("\\u83b7\\u53d6\\u7ae0\\u8282\\u8be6\\u60c5"), "studentId, lessonId, sectionId, time, enc", zh("\\u7ae0\\u8282\\u5185\\u5bb9\\u3001\\u9875\\u9762\\u5217\\u8868\\u3001\\u8fdb\\u5ea6\\u7b49")],
            ["/student-api/api/v1/lesson/resume", "POST", zh("\\u83b7\\u53d6\\u7ee7\\u7eed\\u5b66\\u4e60\\u843d\\u70b9"), "studentId, lessonId, anchorId, anchorTitle, pageNo, time, enc", "anchorId, anchorTitle, pageNo"],
            ["/student-api/api/v1/qa/voiceToText", "POST", zh("\\u8bed\\u97f3\\u8f6c\\u6587\\u5b57"), zh("\\u97f3\\u9891\\u76f8\\u5173 payload\\uff0c\\u5916\\u52a0 time, enc"), zh("\\u8bc6\\u522b\\u6587\\u672c")],
            ["/student-api/ws/qa/asr", "WebSocket", zh("\\u5b9e\\u65f6\\u8bed\\u97f3\\u8bc6\\u522b"), "start / audio " + zh("\\u6d88\\u606f"), "ready, transcript, error"],
            ["/student-api/api/v1/qa/interact", "POST", zh("\\u540c\\u6b65\\u95ee\\u7b54"), "studentId, lessonId, sectionId, question, anchorId, anchorTitle, pageNo, time, enc", "answer, relatedKnowledgePoints, understandingLevel, resumeAnchor, weakPoints"],
            ["/student-api/api/v1/qa/interact/stream", "POST", zh("\\u6d41\\u5f0f\\u95ee\\u7b54 SSE"), zh("\\u540c\\u4e0a\\uff0c\\u524d\\u7aef\\u8fd8\\u4f1a\\u5e26 source, context"), "SSE " + zh("\\u4e8b\\u4ef6") + "\\uff1astart, delta, done, error"],
            ["/student-api/api/v1/qa/sessions/list", "POST", zh("\\u83b7\\u53d6\\u95ee\\u7b54\\u4f1a\\u8bdd\\u5217\\u8868"), "studentId, lessonId, time, enc", "sessions"],
            ["/student-api/api/v1/qa/sessions/save", "POST", zh("\\u4fdd\\u5b58\\u95ee\\u7b54\\u4f1a\\u8bdd"), "studentId, lessonId, sectionId, session, time, enc", zh("\\u4fdd\\u5b58\\u540e\\u7684\\u4f1a\\u8bdd\\u6570\\u636e")],
            ["/student-api/api/v1/notifications/list", "POST", zh("\\u83b7\\u53d6\\u901a\\u77e5\\u5217\\u8868"), "studentId, time, enc", "notifications"],
            ["/student-api/api/v1/notifications/detail", "POST", zh("\\u83b7\\u53d6\\u901a\\u77e5\\u8be6\\u60c5"), "studentId, notificationId, time, enc", zh("\\u901a\\u77e5\\u8be6\\u60c5\\u5bf9\\u8c61")],
            ["/student-api/api/v1/notifications/read", "POST", zh("\\u6807\\u8bb0\\u901a\\u77e5\\u5df2\\u8bfb"), "studentId, notificationId, time, enc", zh("\\u5df2\\u8bfb\\u7ed3\\u679c")],
        ],
    )

    doc.add_heading(zh("\\u95ee\\u7b54\\u76f8\\u5173\\u8fd4\\u56de\\u5b57\\u6bb5"), level=1)
    add_table(
        doc,
        [zh("\\u63a5\\u53e3"), zh("\\u5b57\\u6bb5"), zh("\\u8bf4\\u660e")],
        [
            ["/student-api/api/v1/qa/interact", "answer", zh("\\u0041\\u0049 \\u56de\\u7b54\\u6b63\\u6587")],
            ["/student-api/api/v1/qa/interact", "relatedKnowledgePoints", zh("\\u76f8\\u5173\\u77e5\\u8bc6\\u70b9\\u5217\\u8868")],
            ["/student-api/api/v1/qa/interact", "understandingLevel", zh("\\u7406\\u89e3\\u7a0b\\u5ea6\\uff0c\\u5e38\\u89c1\\u503c\\uff1aweak / partial / complete")],
            ["/student-api/api/v1/qa/interact", "understandingLabel", zh("\\u7406\\u89e3\\u7a0b\\u5ea6\\u4e2d\\u6587\\u6807\\u7b7e")],
            ["/student-api/api/v1/qa/interact", "resumeAnchor", zh("\\u63a8\\u8350\\u7ee7\\u7eed\\u5b66\\u4e60\\u4f4d\\u7f6e")],
            ["/student-api/api/v1/qa/interact", "weakPoints", zh("\\u8584\\u5f31\\u70b9\\u5217\\u8868")],
            ["/student-api/api/v1/qa/interact/stream", "start", zh("\\u6d41\\u5f00\\u59cb\\u4e8b\\u4ef6")],
            ["/student-api/api/v1/qa/interact/stream", "delta", zh("\\u589e\\u91cf\\u6587\\u672c\\u7247\\u6bb5")],
            ["/student-api/api/v1/qa/interact/stream", "done", zh("\\u6700\\u7ec8\\u5b8c\\u6574\\u7ed3\\u679c")],
            ["/student-api/api/v1/qa/interact/stream", "error", zh("\\u9519\\u8bef\\u4e8b\\u4ef6")],
        ],
    )

    doc.add_heading("WebSocket " + zh("\\u5b9e\\u65f6\\u8bed\\u97f3\\u8bc6\\u522b"), level=1)
    add_table(
        doc,
        [zh("\\u63a5\\u53e3"), zh("\\u65b9\\u5411"), zh("\\u6d88\\u606f\\u7c7b\\u578b"), zh("\\u8bf4\\u660e")],
        [
            ["/student-api/ws/qa/asr", zh("\\u5ba2\\u6237\\u7aef -> \\u670d\\u52a1\\u7aef"), "start", zh("\\u521d\\u59cb\\u5316\\u8bc6\\u522b\\uff0c\\u4f1a\\u5e26 studentId, lessonId, sectionId")],
            ["/student-api/ws/qa/asr", zh("\\u5ba2\\u6237\\u7aef -> \\u670d\\u52a1\\u7aef"), "audio", zh("\\u4e0a\\u4f20\\u97f3\\u9891\\u5206\\u7247\\uff0c\\u4f1a\\u5e26 seq, audioBase64, fileName, final")],
            ["/student-api/ws/qa/asr", zh("\\u670d\\u52a1\\u7aef -> \\u5ba2\\u6237\\u7aef"), "ready", zh("\\u670d\\u52a1\\u7aef\\u51c6\\u5907\\u5c31\\u7eea")],
            ["/student-api/ws/qa/asr", zh("\\u670d\\u52a1\\u7aef -> \\u5ba2\\u6237\\u7aef"), "transcript", zh("\\u8fd4\\u56de\\u8bc6\\u522b\\u6587\\u672c\\uff0c\\u53ef\\u80fd\\u5e26 final")],
            ["/student-api/ws/qa/asr", zh("\\u670d\\u52a1\\u7aef -> \\u5ba2\\u6237\\u7aef"), "error", zh("\\u8fd4\\u56de\\u9519\\u8bef\\u4fe1\\u606f")],
        ],
    )

    doc.add_heading(zh("\\u9879\\u76ee\\u5185\\u4fdd\\u7559\\u7684\\u65e7\\u63a5\\u53e3"), level=1)
    add_table(
        doc,
        [zh("\\u63a5\\u53e3"), zh("\\u65b9\\u6cd5"), zh("\\u4f5c\\u7528"), zh("\\u4e3b\\u8981\\u8bf7\\u6c42\\u53c2\\u6570"), zh("\\u5907\\u6ce8")],
        [
            ["/api/v1/qa/voiceToText", "POST", zh("\\u65e7\\u7248\\u8bed\\u97f3\\u8f6c\\u6587\\u5b57"), "voiceUrl, voiceDuration, language, enc, time", zh("\\u65e7\\u94fe\\u8def")],
            ["/api/v1/qa/interact", "POST", zh("\\u65e7\\u7248\\u95ee\\u7b54"), "schoolId, userId, courseId, lessonId, sessionId, questionType, questionContent, currentSectionId, historyQa, enc, time", zh("\\u65e7\\u94fe\\u8def")],
            ["/api/v1/qa/session/{sessionId}", "GET", zh("\\u83b7\\u53d6\\u65e7\\u7248\\u95ee\\u7b54\\u4f1a\\u8bdd"), zh("\\u8def\\u5f84\\u53c2\\u6570 sessionId"), zh("\\u65e7\\u94fe\\u8def")],
        ],
    )

    doc.save(output_path)


if __name__ == "__main__":
    build_document(Path.cwd() / "student-api-interface-docs.docx")
